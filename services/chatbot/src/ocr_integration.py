"""
OCR Integration for Chatbot
Provides text extraction from images and documents
"""

import os
import base64
import requests
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
import io

logger = logging.getLogger(__name__)

# OCR Service URL
OCR_SERVICE_URL = os.getenv("OCR_SERVICE_URL", "http://localhost:5004")


class OCRIntegration:
    """Handles OCR processing for uploaded files"""
    
    def __init__(self, service_url: str = None):
        self.service_url = service_url or OCR_SERVICE_URL
        self.enabled = True
        self._check_service()
    
    def _check_service(self):
        """Check if OCR service is available"""
        try:
            response = requests.get(f"{self.service_url}/health", timeout=2)
            self.enabled = response.status_code == 200
            if self.enabled:
                logger.info("✅ OCR Service connected")
            else:
                logger.warning("⚠️ OCR Service unavailable - using fallback")
        except:
            self.enabled = False
            logger.warning("⚠️ OCR Service not reachable - using fallback")
    
    def extract_text_from_image(self, image_data: bytes, filename: str = "image.png") -> Dict[str, Any]:
        """
        Extract text from image using OCR
        
        Args:
            image_data: Raw image bytes
            filename: Original filename
            
        Returns:
            Dict with extracted text and metadata
        """
        result = {
            "success": False,
            "text": "",
            "confidence": 0,
            "language": "unknown",
            "method": "none"
        }
        
        # Try external OCR service first
        if self.enabled:
            try:
                files = {"file": (filename, image_data)}
                response = requests.post(
                    f"{self.service_url}/api/ocr",
                    files=files,
                    timeout=30
                )
                
                if response.status_code == 200:
                    data = response.json()
                    result["success"] = True
                    result["text"] = data.get("text", "")
                    result["confidence"] = data.get("confidence", 0.9)
                    result["language"] = data.get("language", "vi")
                    result["method"] = "paddle_ocr"
                    logger.info(f"[OCR] Extracted {len(result['text'])} chars from {filename}")
                    return result
            except Exception as e:
                logger.warning(f"[OCR] Service error: {e}")
        
        # Fallback: Try using OpenAI Vision API
        result = self._fallback_vision_ocr(image_data, filename)
        return result
    
    def _fallback_vision_ocr(self, image_data: bytes, filename: str) -> Dict[str, Any]:
        """Use OpenAI Vision API as fallback OCR"""
        result = {
            "success": False,
            "text": "",
            "confidence": 0,
            "language": "unknown",
            "method": "vision_api"
        }
        
        try:
            import openai
            api_key = os.getenv("OPENAI_API_KEY")
            
            if not api_key:
                logger.warning("[OCR] No OpenAI API key for fallback OCR")
                return result
            
            # Convert image to base64
            base64_image = base64.b64encode(image_data).decode('utf-8')
            
            # Detect image type
            if filename.lower().endswith('.png'):
                media_type = "image/png"
            elif filename.lower().endswith(('.jpg', '.jpeg')):
                media_type = "image/jpeg"
            elif filename.lower().endswith('.gif'):
                media_type = "image/gif"
            elif filename.lower().endswith('.webp'):
                media_type = "image/webp"
            else:
                media_type = "image/png"
            
            client = openai.OpenAI(api_key=api_key)
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extract ALL text from this image. Return ONLY the extracted text, nothing else. If there's no text, return 'NO_TEXT_FOUND'."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{media_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000
            )
            
            extracted = response.choices[0].message.content.strip()
            
            if extracted and extracted != "NO_TEXT_FOUND":
                result["success"] = True
                result["text"] = extracted
                result["confidence"] = 0.85
                result["language"] = "auto"
                logger.info(f"[OCR] Vision API extracted {len(extracted)} chars from {filename}")
            
            return result
            
        except Exception as e:
            logger.error(f"[OCR] Vision API fallback failed: {e}")
            return result
    
    def extract_text_from_pdf(self, pdf_data: bytes, filename: str = "document.pdf") -> Dict[str, Any]:
        """Extract text from PDF"""
        result = {
            "success": False,
            "text": "",
            "pages": 0,
            "method": "none"
        }
        
        # Try external service
        if self.enabled:
            try:
                files = {"file": (filename, pdf_data)}
                response = requests.post(
                    f"{self.service_url}/api/ocr",
                    files=files,
                    timeout=60
                )
                
                if response.status_code == 200:
                    data = response.json()
                    result["success"] = True
                    result["text"] = data.get("text", "")
                    result["pages"] = data.get("pages", 1)
                    result["method"] = "ocr_service"
                    return result
            except Exception as e:
                logger.warning(f"[OCR] PDF service error: {e}")
        
        # Fallback: Use PyPDF2
        try:
            import PyPDF2
            pdf_file = io.BytesIO(pdf_data)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            
            result["success"] = True
            result["text"] = "\n\n".join(text_parts)
            result["pages"] = len(reader.pages)
            result["method"] = "pypdf2"
            logger.info(f"[OCR] Extracted {len(result['text'])} chars from {filename}")
            
        except Exception as e:
            logger.error(f"[OCR] PDF extraction failed: {e}")
        
        return result
    
    def process_file(self, file_data: bytes, filename: str, content_type: str = None) -> Dict[str, Any]:
        """
        Process any file and extract text
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            content_type: MIME type (optional)
            
        Returns:
            Dict with extracted content
        """
        ext = Path(filename).suffix.lower()
        
        # Image files
        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tiff']:
            return self.extract_text_from_image(file_data, filename)
        
        # PDF files
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_data, filename)
        
        # Text-based files - just decode
        if ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv']:
            try:
                text = file_data.decode('utf-8')
                return {
                    "success": True,
                    "text": text,
                    "method": "direct_read",
                    "file_type": ext[1:]
                }
            except:
                try:
                    text = file_data.decode('latin-1')
                    return {
                        "success": True,
                        "text": text,
                        "method": "direct_read",
                        "file_type": ext[1:]
                    }
                except:
                    pass
        
        # Word documents
        if ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_data, filename)
        
        # Excel files
        if ext in ['.xlsx', '.xls']:
            return self._extract_from_excel(file_data, filename)
        
        return {
            "success": False,
            "text": "",
            "error": f"Unsupported file type: {ext}"
        }
    
    def _extract_from_docx(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from Word documents"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_data))
            
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(row_text))
            
            return {
                "success": True,
                "text": "\n".join(text_parts),
                "method": "python_docx",
                "file_type": "docx"
            }
        except Exception as e:
            logger.error(f"[OCR] DOCX extraction failed: {e}")
            return {"success": False, "text": "", "error": str(e)}
    
    def _extract_from_excel(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from Excel files"""
        try:
            import pandas as pd
            
            xlsx = pd.ExcelFile(io.BytesIO(file_data))
            
            text_parts = []
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                text_parts.append(df.to_string())
            
            return {
                "success": True,
                "text": "\n\n".join(text_parts),
                "method": "pandas",
                "file_type": "xlsx",
                "sheets": xlsx.sheet_names
            }
        except Exception as e:
            logger.error(f"[OCR] Excel extraction failed: {e}")
            return {"success": False, "text": "", "error": str(e)}


# Global instance
ocr_client = OCRIntegration()


def extract_file_content(file_data: bytes, filename: str) -> Tuple[bool, str]:
    """
    Convenience function to extract content from file
    
    Returns:
        Tuple of (success, extracted_text)
    """
    result = ocr_client.process_file(file_data, filename)
    return result.get("success", False), result.get("text", "")
